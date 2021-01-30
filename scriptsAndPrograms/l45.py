import docx
d = docx.Document('..\\textFile\\demo.docx')
print(d.paragraphs)
print((d.paragraphs[0]).text)

p = d.paragraphs[1]
print(p.runs)
print(p.runs[0].text)
print(p.runs[1].text)
print(p.runs[2].text)
print(p.runs[3].text)

print(p.runs[1].text)

print(p.runs[1].bold)
p.runs[1].bold == None
print(p.runs[1].bold)

print(p.runs[3].italic)
p.runs[3].underline = True
p.runs[3].text = 'italic and underline'

d.save('..\\textFile\\demo2.docx')
p.style
p.style = 'Title'
d.save('..\\textFile\\demo3.docx')

d = docx.Document()
d.add_paragraph('Hello this is a paragraph.')
d.add_paragraph('Hello this is a new paragraph.')
d.save('..\\textFile\\demo4.docx')

p = d.paragraphs[0]
p.add_run('This is a new run')

p.runs
p.runs[1].bold = True
d.save('..\\textFile\\demo5.docx')

